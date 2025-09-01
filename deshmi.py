# file: deshmi_penaliteti_app.py
import os, re, zipfile, tempfile, shlex, subprocess
from io import BytesIO
from datetime import datetime
from typing import Dict, Any, List, Tuple

# single, global qn alias (use this everywhere)
from docx.oxml.ns import qn as _qn
from docx.oxml import OxmlElement


# ── Streamlit must be configured before any other st.* call
import streamlit as st
st.set_page_config(page_title="Deshmi Penaliteti", layout="centered")

# ── Env
from dotenv import load_dotenv
load_dotenv()

AWS_ACCESS_KEY_ID     = os.getenv("AWS_ACCESS_KEY_ID")
AWS_SECRET_ACCESS_KEY = os.getenv("AWS_SECRET_ACCESS_KEY")
AWS_REGION            = os.getenv("AWS_REGION", "us-east-2")
APP_PASSWORD          = os.getenv("APP_PASSWORD")  # optional locally

import boto3
textract = boto3.client(
    "textract",
    aws_access_key_id     = AWS_ACCESS_KEY_ID,
    aws_secret_access_key = AWS_SECRET_ACCESS_KEY,
    region_name           = AWS_REGION,
)


# ── DOCX stuff
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Mm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL

# ────────────────────────────────────────────────────────────────────────────
# UI
# ────────────────────────────────────────────────────────────────────────────
st.title("Vërtetim i Gjendjes Gjyqësore → Shqip-Italisht")

if APP_PASSWORD:
    pw = st.text_input("Password", type="password")
    if pw != APP_PASSWORD:
        st.warning("Fut password-in për të vazhduar.")
        st.stop()
else:
    st.info("APP_PASSWORD nuk është vendosur në .env. Duke ekzekutuar pa fjalëkalim (lokalisht).")

uploaded_files = st.file_uploader(
    "Ngarko dokumente (PDF/JPG/PNG)", type=["pdf", "jpg", "jpeg", "png"], accept_multiple_files=True
)
download_format = st.selectbox("Formati i daljes", ["Word (.docx)", "PDF (.pdf)"])

# ────────────────────────────────────────────────────────────────────────────
# Textract helpers
# ────────────────────────────────────────────────────────────────────────────
def run_textract(file_bytes: bytes) -> Dict[str, Any]:
    return textract.analyze_document(
        Document={'Bytes': file_bytes},
        FeatureTypes=["FORMS", "TABLES", "LAYOUT"]
    )

def blocks_map(resp: Dict[str, Any]) -> Tuple[List[Dict[str,Any]], Dict[str,Dict[str,Any]]]:
    blocks = resp["Blocks"]
    bmap = {b["Id"]: b for b in blocks}
    return blocks, bmap

def all_lines(blocks: List[Dict[str,Any]]) -> List[Dict[str,Any]]:
    return [b for b in blocks if b["BlockType"] == "LINE" and b.get("Text")]

def deaccent_e(text: str) -> str:
    return text.replace("ë", "e").replace("Ë", "E")

def y_center(bb) -> float:
    return bb["Top"] + bb["Height"]/2.0

def same_line_y(bb1, bb2, tol=0.015) -> bool:
    return abs(y_center(bb1) - y_center(bb2)) <= tol

def nearest_right_value(blocks, label_line, prefer_regex: str = None):
    bb = label_line["Geometry"]["BoundingBox"]
    cands = []
    for ln in all_lines(blocks):
        if ln["Id"] == label_line["Id"]:
            continue
        bb2 = ln["Geometry"]["BoundingBox"]
        if same_line_y(bb, bb2) and bb2["Left"] > bb["Left"]:
            txt = ln["Text"].strip()
            if not txt:
                continue
            if prefer_regex:
                if re.search(prefer_regex, txt):
                    cands.append((bb2["Left"], txt))
            else:
                cands.append((bb2["Left"], txt))
    if not cands:
        return ""
    cands.sort(key=lambda t: t[0])
    return cands[0][1]

# Canonical exonyms (Italian)
_EXO = {
    "tirane": "Tirana",
    "tiranë": "Tirana",
    "durres": "Durazzo",
    "durrës": "Durazzo",
    "vlore":  "Valona",
    "vlorë":  "Valona",
}

def _deacc(s: str) -> str:
    return (s or "").replace("ë","e").replace("Ë","E")\
                    .replace("ç","c").replace("Ç","C")

def _match_case(target: str, src: str) -> str:
    """Make `target` follow the casing style of `src`."""
    if src.isupper():
        return target.upper()
    if src.islower():
        return target.lower()
    if src[:1].isupper() and src[1:].islower():
        return target.capitalize()
    # Mixed/unknown casing -> leave target as-is
    return target

def normalize_city(city: str) -> str:
    """Normalize a single city field to Italian exonym while keeping casing."""
    if not city:
        return ""
    s = city.strip()
    key = s.lower()
    ex = _EXO.get(key) or _EXO.get(_deacc(key))
    return _match_case(ex, s) if ex else s

def normalize_cities_in_text(text: str) -> str:
    """Replace ALL occurrences in free text, preserving each token's casing."""
    if not text:
        return text

    # All variants we want to catch (accented + unaccented)
    variants = ["tiranë","tirane","durrës","durres","vlorë","vlore"]
    pattern = r"\b(" + "|".join(map(re.escape, variants)) + r")\b"

    def repl(m):
        found = m.group(0)
        ex = _EXO.get(found.lower()) or _EXO.get(_deacc(found.lower()))
        return _match_case(ex, found) if ex else found

    return re.sub(pattern, repl, text, flags=re.IGNORECASE)

def parse_name_surname_line(txt: str) -> tuple[str, str]:
    s = re.sub(r"[,\u200b]+", " ", txt or "").strip()
    s = re.sub(r"\s+", " ", s)
    m = re.match(r"^([A-ZÇË' -]+?)\s+([A-ZÇË' -]+?)(\s*\([A-ZÇË' -]+\))?\s*$", s)
    if m:
        name = m.group(1).strip()
        surname = (m.group(2) + (m.group(3) or "")).strip()
        return name, surname
    if not s:
        return "", ""
    parts = s.split()
    name = parts[0]
    base_surname = s.split("(")[0].split()[-1] if "(" in s else (parts[-1] if len(parts) > 1 else "")
    paren = re.search(r"(\s*\([A-ZÇË' -]+\))\s*$", s)
    surname = (base_surname + (paren.group(1) if paren else "")).strip()
    return name, surname

def set_cell_top_border(cell, size="8", color="000000"):
    tc_pr = cell._tc.get_or_add_tcPr()

    # ensure a <w:tcBorders> exists
    tc_borders = tc_pr.find(_qn('w:tcBorders'))
    if tc_borders is None:
        tc_borders = OxmlElement('w:tcBorders')
        tc_pr.append(tc_borders)

    # ensure a <w:top> exists and set attrs
    top = tc_borders.find(_qn('w:top'))
    if top is None:
        top = OxmlElement('w:top')
        tc_borders.append(top)

    top.set(_qn('w:val'),   'single')
    top.set(_qn('w:sz'),    size)      # thickness
    top.set(_qn('w:space'), '0')
    top.set(_qn('w:color'), color)

def remove_table_borders(table):
    """Remove all borders from a python-docx table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl._tbl.append(tblPr)
    borders = OxmlElement('w:tblBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(_qn('w:val'), 'nil')     # no border
        borders.append(el)
    # drop any existing borders and add ours
    for old in tblPr.findall(_qn('w:tblBorders')):
        tblPr.remove(old)
    tblPr.append(borders)

def add_kv_table(doc, rows, left_w_cm=7.5, right_w_cm=9.0, font_size_pt=14):
    """
    rows: list[tuple[label, value or list of (text, bold?) runs]]
    Allows mixed bold/normal runs in the right column.
    """
    tbl = doc.add_table(rows=0, cols=2)
    tbl.autofit = False
    tbl.columns[0].width = Cm(left_w_cm)
    tbl.columns[1].width = Cm(right_w_cm)

    for label, value in rows:
        r = tbl.add_row()
        c1, c2 = r.cells
        c1.width = Cm(left_w_cm); c2.width = Cm(right_w_cm)

        # left label always bold
        p1 = c1.paragraphs[0]; p1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run1 = p1.add_run(str(label))
        run1.font.name = 'Times New Roman'
        run1.font.size = Pt(font_size_pt)
        run1.bold = False

        # right column, mixed runs
        p2 = c2.paragraphs[0]; p2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        if isinstance(value, list):
            # if you pass [(text, bold), ...]
            for txt, is_bold in value:
                run2 = p2.add_run(txt)
                run2.font.name = 'Times New Roman'
                run2.font.size = Pt(14)
                run2.bold = bool(is_bold)
        else:
            run2 = p2.add_run(str(value) if value is not None else "")
            run2.font.name = 'Times New Roman'
            run2.font.size = Pt(font_size_pt)
            run2.bold = True   # default bold if not list

    remove_table_borders(tbl)
    return tbl


# ---- compact signer extraction --------------------------------------------
def _is_name_like(s: str) -> bool:
    s = s.strip()
    if len(s.split()) < 2:
        return False
    # "Metvaldo Hiraj" (with ë/ç allowed)
    if re.fullmatch(r"[A-ZÇË][a-zçë]+(?:\s+[A-ZÇË][a-zçë]+)+", s):
        return True
    # "METVALDO HIRAJ"
    if re.fullmatch(r"[A-ZÇË]+(?:\s+[A-ZÇË]+)+", s):
        return True
    return False

def extract_signer_from_lines(lines):
    # find "Sektori i Gjendjes Gjyqësore"
    idx = next(
        (i for i, ln in enumerate(lines)
         if "sektori" in deaccent_e(ln["Text"].lower())
         and "gjendjes" in deaccent_e(ln["Text"].lower())),
        None
    )
    # 1) try the next few lines after the anchor
    if idx is not None:
        for ln in lines[idx+1 : idx+8]:
            t = (ln["Text"] or "").strip()
            if _is_name_like(t):
                return t
    # 2) fallback: bottom of page
    for ln in reversed(lines[-12:]):
        t = (ln["Text"] or "").strip()
        if _is_name_like(t):
            return t
    return ""
# ---------------------------------------------------------------------------

# ── HELPER: SEAL FOOTER (robust, anchor on Vulosur) ─────────────────────────
# ── HELPER: E-SEAL (flex length ≥ 20 hex) ───────────────────────────────────
def extract_seal_footer(blocks, which="last", min_len=20):
    """
    Extract the electronic-seal footer near 'Vulosur elektronikisht'.
    Returns a 4-line Italian block or "".

    which: "first" | "second" | "last"
    min_len: minimum hex length for the seal id (default 20)
    """
    import re

    def deacc(s: str) -> str:
        return (s or "").lower().replace("ë", "e")

    # Collect LINEs (keep geometry for band fallback)
    lines = [b for b in blocks if b.get("BlockType") == "LINE" and b.get("Text")]
    if not lines:
        return ""

    # Find anchor index(es)
    hits = [i for i, ln in enumerate(lines) if "vulosur elektronikisht" in deacc(ln["Text"])]
    if not hits:
        return ""

    if which == "first":
        start = hits[0]
    elif which == "second" and len(hits) >= 2:
        start = hits[1]
    else:
        start = hits[-1]

    # Small window of lines after the anchor
    tail = lines[start : min(len(lines), start + 12)]
    snippet = "\n".join((ln["Text"] or "").strip() for ln in tail)

    # 1) Date (yyyy/mm/dd with optional time / tz)
    date_line = ""
    date_re = re.compile(
        r"\b\d{4}/\d{2}/\d{2}"             # YYYY/MM/DD
        r"(?:[ T]\d{2}:\d{2}:\d{2}"        # HH:MM:SS
        r"(?:\s*[+-]\d{2}[:'’]?\d{2})?)?", # timezone
        re.UNICODE
    )
    m = date_re.search(snippet)
    if m:
        cleaned = re.sub(r"^(Date|Datë|Daté)\s*:?\s*", "", m.group(0), flags=re.I).strip()
        date_line = f"In data {cleaned}"

    # 2) Hex seal (any contiguous hex run ≥ min_len)
    hex_re = re.compile(rf"\b[0-9a-fA-F]{{{min_len},}}\b")
    candidates = hex_re.findall(snippet)
    hash_line = max(candidates, key=len) if candidates else ""

    # Fallback: scan WORDs in the same vertical band as the anchor
    if not hash_line:
        words = [b for b in blocks if b.get("BlockType") == "WORD" and b.get("Text")]
        abb = lines[start]["Geometry"]["BoundingBox"]
        y0, y1 = max(0.0, abb["Top"] - 0.03), min(1.0, abb["Top"] + 0.25)
        x0, x1 = 0.10, 0.98  # skip far-left QR zone
        band_text = " ".join(
            (w["Text"] or "").strip()
            for w in words
            if x0 <= (w["Geometry"]["BoundingBox"]["Left"] + w["Geometry"]["BoundingBox"]["Width"]/2) <= x1
            and y0 <= (w["Geometry"]["BoundingBox"]["Top"]  + w["Geometry"]["BoundingBox"]["Height"]/2) <= y1
        )
        candidates = hex_re.findall(band_text)
        hash_line = max(candidates, key=len) if candidates else ""

    if not (date_line or hash_line):
        return ""

    return "\n".join([
        "Timbrato elettronicamente dalla",
        "Direzione Generale delle Carceri",
        date_line,
        hash_line
    ])



# ────────────────────────────────────────────────────────────────────────────
# Field extraction
# ────────────────────────────────────────────────────────────────────────────
def extract_fields(blocks: List[Dict[str,Any]]) -> Dict[str,str]:
    lines = all_lines(blocks)
    T = "\n".join(ln["Text"] for ln in lines)
    out = {
        "request_no": "", "city": "", "request_date": "",
        "name": "", "surname": "",
        "father_name": "", "mother_name": "",
        "dob": "", "birthplace": "", "personal_no": "",
        "status_text": "", "signer": "",
        "e_seal": "",  # ← add this
    }

    for ln in lines:
        t = deaccent_e(ln["Text"].lower())
        if "nr" in t and "kerkese" in t:
            out["request_no"] = nearest_right_value(blocks, ln, r"[A-Za-z0-9/.-]+").strip()
            break

    for ln in lines:
        raw = ln["Text"]
        low = deaccent_e(raw)
        if re.search(r"\bm[eë]\b", low, flags=re.I):
            parts = re.split(r"\bm[eë]\b", raw, flags=re.I)
            if parts and parts[0].strip():
                out["city"] = parts[0].strip()
                rd = nearest_right_value(blocks, ln, r"\d{2}/\d{2}/\d{4}")
                if rd:
                    m = re.search(r"\d{2}/\d{2}/\d{4}", rd)
                    if m: out["request_date"] = m.group(0)
                break

    # ---------- Name / Surname (inline label; geometry-based) ----------
    name, surname = "", ""

    # 1) find the LINE that contains the label
    label_line = next(
        (ln for ln in lines
        if "emri" in deaccent_e(ln["Text"].lower())
        and "mbiemri" in deaccent_e(ln["Text"].lower())),
        None
    )

    if label_line:
        line_bb = label_line["Geometry"]["BoundingBox"]

        # 2) find the right edge of the label (“mbiemri” and any trailing “)” word)
        cut_x = None
        last_right_edge_after_mbiemri = None
        mbiemri_seen = False

        for w in blocks:
            if w.get("BlockType") != "WORD" or not w.get("Text"):
                continue
            wbb = w["Geometry"]["BoundingBox"]
            if not same_line_y(line_bb, wbb, tol=0.02):
                continue

            wt = deaccent_e(w["Text"].lower())
            right_edge = wbb["Left"] + wbb["Width"]

            if "mbiemri" in wt:
                mbiemri_seen = True
                last_right_edge_after_mbiemri = right_edge
                continue

            # if Textract split the trailing “)” into its own word, extend the edge
            if mbiemri_seen and wt.strip() in (")", ").", "),"):
                last_right_edge_after_mbiemri = max(last_right_edge_after_mbiemri or right_edge, right_edge)

        if last_right_edge_after_mbiemri is not None:
            cut_x = last_right_edge_after_mbiemri

        # 3) collect VALUE words strictly to the right of the label
        value_words = []
        if cut_x is not None:
            for w in blocks:
                if w.get("BlockType") != "WORD" or not w.get("Text"):
                    continue
                wbb = w["Geometry"]["BoundingBox"]
                if same_line_y(line_bb, wbb, tol=0.02) and (wbb["Left"] > cut_x + 0.002):
                    value_words.append((wbb["Left"], w["Text"]))

        value_words.sort(key=lambda t: t[0])
        tokens = [t for _, t in value_words]

        # 4) if nothing captured (rare scan), fall back to "next line" words
        if not tokens:
            # find index of the label line
            idx = next((i for i, ln in enumerate(lines) if ln is label_line), None)
            if idx is not None and idx + 1 < len(lines):
                next_line = lines[idx + 1]
                nbb = next_line["Geometry"]["BoundingBox"]
                band = sorted(
                    [(w["Geometry"]["BoundingBox"]["Left"], w["Text"])
                    for w in blocks
                    if w.get("BlockType") == "WORD" and w.get("Text")
                    and same_line_y(nbb, w["Geometry"]["BoundingBox"], tol=0.02)],
                    key=lambda t: t[0]
                )
                tokens = [t for _, t in band]

        # 5) build name + surname from tokens (KEEP parentheses)
        if tokens:
            name = tokens[0].strip()
            surname = " ".join(tokens[1:]).strip()

    out["name"] = name
    out["surname"] = surname
    # ------------------------------------------------------



    for ln in lines:
        t = deaccent_e(ln["Text"].lower())
        if ("i biri" in t or "e bija" in t) and t.strip().endswith("i"):
            out["father_name"] = (nearest_right_value(blocks, ln, r"[A-ZÇË][A-Za-zÇËçë\-() ]+") or "").strip()
            break

    for ln in lines:
        t = deaccent_e(ln["Text"].lower().strip())
        if t == "dhe i":
            out["mother_name"] = (nearest_right_value(blocks, ln, r"[A-ZÇË][A-Za-zÇËçë\-() ]+") or "").strip()
            break

    m = re.search(r"lindur\s+m[ëe]\s+(\d{2}/\d{2}/\d{4}).{0,30}n[ëe]\s+([A-ZÇË ,.-]+)", T, flags=re.I|re.S)
    if m:
        out["dob"] = m.group(1).strip()
        out["birthplace"] = m.group(2).strip().replace("\n"," ").replace(" ,", ",")
    else:
        for ln in lines:
            if re.search(r"lindur\s+m[ëe]\b", deaccent_e(ln["Text"]), flags=re.I):
                d = nearest_right_value(blocks, ln, r"\d{2}/\d{2}/\d{4}")
                if d:
                    out["dob"] = re.search(r"\d{2}/\d{2}/\d{4}", d).group(0)
                after = nearest_right_value(blocks, ln, r".+")
                if after:
                    m2 = re.search(r"n[ëe]\s+(.+)$", deaccent_e(after), flags=re.I)
                    if m2: out["birthplace"] = m2.group(1).strip()
                break

    for ln in lines:
        t = deaccent_e(ln["Text"].lower())
        if "me numer personal" in t:
            out["personal_no"] = (nearest_right_value(blocks, ln, r"[A-Za-z0-9]+") or "").strip()
            break

    # Always normalize to Italian wording
    for ln in lines:
        t = ln["Text"].strip()
        if re.search(r"pa\s+d[ëe]nuar", deaccent_e(t), flags=re.I):
            out["status_text"] = "RISULTA INCENSURATO"
            break
    if not out["status_text"]:
        out["status_text"] = "RISULTA INCENSURATO"

    signer = extract_signer_from_lines(lines)
    out["signer"] = signer

    out["e_seal"] = extract_seal_footer(blocks, which="first")

    out["city"] = normalize_city(out["city"].strip().split(",")[0])
    out["birthplace"] = normalize_cities_in_text(out.get("birthplace", ""))

    return out

# ────────────────────────────────────────────────────────────────────────────
# DOCX builder
# ────────────────────────────────────────────────────────────────────────────
def add_p(doc, text, size=11, bold=False, align="left", italic=False, indent_cm=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0,0,0)
    run.bold = bold
    run.italic = italic
    p.alignment = {
        "left":   WD_PARAGRAPH_ALIGNMENT.LEFT,
        "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
        "right":  WD_PARAGRAPH_ALIGNMENT.RIGHT,
        "justify":WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    }.get(align, WD_PARAGRAPH_ALIGNMENT.LEFT)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
        # add indentation if requested
    if indent_cm is not None:
        p.paragraph_format.left_indent = Cm(indent_cm)
    return p

def build_docx(data: Dict[str,str]) -> BytesIO:
    doc = Document()

    today = datetime.today().strftime("%d.%m.%Y")
    section = doc.sections[0]
    section.top_margin    = Cm(1.7)
    section.bottom_margin = Cm(1.3)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)
    section.page_width  = Mm(210)
    section.page_height = Mm(297)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    # <<< use the global alias, and no inner import anywhere >>>
    style.element.rPr.rFonts.set(_qn('w:eastAsia'), 'Times New Roman')

    flag_path = os.path.join(os.getcwd(), "al_flag.png")  # or al_coat.png
    if os.path.exists(flag_path):
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        r = p.add_run()
        r.add_picture(flag_path, width=Cm(1.3))  # adjust width

    # Ministry block in body (not header)
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run(
        "REPUBBLICA D’ALBANIA\n"
        "MINISTERO DI GIUSTIZIA\n"
        "Direzione Generale delle Carceri"
    )
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0,0,0)

    # Address + Tel/Fax line (with top border)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = False
    tbl.columns[0].width = Cm(9)
    tbl.columns[1].width = Cm(7.5)

    # apply top border to BOTH cells
    left_cell, right_cell = tbl.rows[0].cells
    set_cell_top_border(left_cell)
    set_cell_top_border(right_cell)

    # left text
    p1 = left_cell.paragraphs[0]
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p1.add_run('Indirizzo: Via “Zef Serembe”').font.size = Pt(12)

    # right text
    p2 = right_cell.paragraphs[0]
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    p2.add_run('Tel/Fax: 00355 4 22 82 92').font.size = Pt(12)

    # Meta row (Nr. Kërkese — City + Date)
    meta_tbl = doc.add_table(rows=1, cols=2)
    meta_tbl.autofit = False
    meta_tbl.columns[0].width = Cm(9)
    meta_tbl.columns[1].width = Cm(7.5)
    c1, c2 = meta_tbl.rows[0].cells

    p1 = c1.paragraphs[0]; p1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    r1 = p1.add_run("Nr. di domanda "); r1.bold = False; r1.underline = True; r1.font.size = Pt(14)
    r2 = p1.add_run(data.get("request_no", "").strip()); r2.underline = True; r2.font.size = Pt(14)

    p2 = c2.paragraphs[0]; p2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    r3 = p2.add_run(f"{data.get('city','Tiranë').strip()} lì "); r3.bold = False; r3.underline = True; r3.font.size = Pt(14)
    r4 = p2.add_run(data.get("request_date", "").strip()); r4.underline = True; r4.font.size = Pt(14)

    add_p(doc, "CERTIFICATO\nDEL CASELLARIO GIUDIZIALE\n", bold=True, align="center", size=16)

    # --- Body replacement: use a 2-column borderless table ---
    preamble = (
        "In applicazione dell’articolo 484 del Codice di Procedura Penale, "
        "della Repubblica d’Albania, dagli accertamenti effettuati sul registro "
        "giudiziario presso questo Ministero risulta che il/la cittadino/a:"
    )
    add_p(doc, preamble, align="justify", size=14)
    doc.add_paragraph()  # spacer

    rows = [
        ("(nome, cognome)", [(f"{data.get('name','')} {data.get('surname','')}", True)]),
    ]

    # Father + Mother
    parents_val = []
    if data.get("father_name"):
        parents_val.append((data.get("father_name").strip(), True))
    if data.get("mother_name"):
        parents_val.append(("  e di  ", False))
        parents_val.append((data['mother_name'].strip(), True))
    rows.append(("figlio (figlia) di", parents_val))

    # Date + birthplace
    dob_val = []
    if data.get("dob"):
        dob_val.append((data.get("dob").strip(), True))
    if data.get("birthplace"):
        dob_val.append(("   a   ", False))
        dob_val.append((data.get("birthplace").strip(), True))
    rows.append(("nato/a il", dob_val))

    # Personal no
    rows.append(("con numero personale", [(data.get("personal_no",""), True)]))

    add_kv_table(doc, rows, left_w_cm=6.0, right_w_cm=11.0, font_size_pt=14)
    doc.add_paragraph()  # spacer


    add_p(doc, "" + data.get("status_text","RISULTA INCENSURATO") + "\n", bold=True, size=16)
    add_p(doc, "Settore di Casellario Giudiziale", align="center", size=14,  indent_cm=6, bold=True)
    if data.get("signer"):
        add_p(doc, f"{data['signer']}", align="center", size=14, indent_cm=6)

    # === ▼ STEP 3: render the e-seal footer lines here ▼ ===
    seal = (data.get("e_seal") or "").strip()
    if seal:
        for line in seal.splitlines():
            add_p(doc, line, size=10, align="left", italic=True)

    add_p(doc,
          "\nAnnotazione: Il presente documento è generato e timbrato\ntramite una procedura automatica dal sistema elettronico\n(Direzione Generale delle Carceri)\n",
          size=10, italic=True)

    table = doc.add_table(rows=1, cols=1)
    table.autofit = False  # Disable Word's auto-resizing
    table.style = 'Table Grid'

    table.columns[0].width = Cm(11)
    table.rows[0].cells[0].width = Cm(11)  # Redundant but safer for compatibility

    cell = table.rows[0].cells[0]
    p = cell.paragraphs[0]  
    run = p.add_run(
        "Io, Vjollca META, traduttrice ufficiale della lingua italiana certificata dal  "
        "Ministero della Giustizia con il numero di certificato 412 datato 31.07.2024, "
        "dichiaro di aver tradotto il testo che mi è stato presentato dalla lingua "
        "albanese nella lingua italiana con precisione, con la dovuta diligenza e "
        "responsabilità legale.\n"
        f"In data {today}."
)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    add_p(
        doc,
        "\nTraduzione eseguita da:\nVjollca META",
        size=11,
        align="center",
        indent_cm=12
    )


    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def docx_to_pdf_bytes(docx_bytes: bytes) -> bytes:
    with tempfile.TemporaryDirectory() as tmp:
        docx_path = os.path.join(tmp, "tmp.docx")
        pdf_path  = os.path.join(tmp, "tmp.pdf")
        with open(docx_path, "wb") as f: f.write(docx_bytes)
        try:
            from docx2pdf import convert
            convert(docx_path, pdf_path)
        except Exception:
            cmd = f'libreoffice --headless --convert-to pdf --outdir "{tmp}" "{docx_path}"'
            subprocess.run(shlex.split(cmd), check=True)
        with open(pdf_path, "rb") as f:
            return f.read()

# ────────────────────────────────────────────────────────────────────────────
# Main
# ────────────────────────────────────────────────────────────────────────────
def process_one(upload) -> Tuple[Dict[str,str], bytes, str]:
    resp = run_textract(upload.read())
    blocks, _ = blocks_map(resp)
    data = extract_fields(blocks)
    doc_buf = build_docx(data)
    if download_format.startswith("PDF"):
        return data, docx_to_pdf_bytes(doc_buf.getvalue()), "pdf"
    else:
        return data, doc_buf.getvalue(), "docx"

if uploaded_files and st.button("✅ Përkthe"):
    if len(uploaded_files) == 1:
        up = uploaded_files[0]
        with st.spinner("Duke nxjerrë fushat dhe duke ndërtuar dokumentin…"):
            data, out_bytes, ext = process_one(up)
        with st.expander("🔎 Fushat e nxjerra"): st.json(data)
        name_part = f'{(data.get("name") or "EMER").strip().replace(" ","_")}_{(data.get("surname") or "MBIEMER").strip().replace(" ","_")}'
        today = datetime.today().strftime("%Y-%m-%d")
        fn = f"{name_part}_Vertetim_Gjyqesor_{today}.{ext}"
        st.download_button("📥 Shkarko", out_bytes, file_name=fn,
                           mime="application/pdf" if ext=="pdf" else "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    else:
        zip_buf = BytesIO()
        with zipfile.ZipFile(zip_buf, "w") as zf:
            for up in uploaded_files:
                with st.spinner(f"Po përpunohet: {up.name}"):
                    data, out_bytes, ext = process_one(up)
                    name_part = f'{(data.get("name") or "EMER").strip().replace(" ","_")}_{(data.get("surname") or "MBIEMER").strip().replace(" ","_")}'
                    today = datetime.today().strftime("%Y-%m-%d")
                    fn = f"{name_part}_Vertetim_Gjyqesor_{today}.{ext}"
                    zf.writestr(fn, out_bytes)
        zip_buf.seek(0)
        st.download_button("📦 Shkarko të gjitha (ZIP)", data=zip_buf,
                           file_name=f"vertetime_{datetime.today().strftime('%Y-%m-%d')}.zip",
                           mime="application/zip")
